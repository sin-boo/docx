import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from output_utils import save_docx_and_pdf

# ── Colour palette ─────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1A, 0x3A, 0x5C)
TEAL  = RGBColor(0x00, 0x7A, 0x87)
GOLD  = RGBColor(0xE6, 0xA8, 0x17)
LIGHT = RGBColor(0xF0, 0xF6, 0xFA)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
GREY  = RGBColor(0x6B, 0x7B, 0x8D)
MINT  = RGBColor(0xD4, 0xF1, 0xF4)


WORKSHEETS = [
    {
        "slug": "hira_01_vowels.docx",
        "title": "Hiragana Set 1 — Vowels",
        "subtitle": "あ い う え お",
        "focus": "Learn the five basic vowel sounds. These vowels appear in almost every Japanese word.",
        "kana": [
            ("あ", "a", "Open and clear, like a in father."),
            ("い", "i", "Short i, like ee in see."),
            ("う", "u", "Short u, with rounded lips."),
            ("え", "e", "Short e, like e in bed."),
            ("お", "o", "Short o, like o in open."),
        ],
        "pairs": [("あい", "a-i"), ("うえ", "u-e"), ("あお", "a-o"), ("いえ", "i-e")],
        "teaching_notes": [
            "These five vowels are the sound base for all the other hiragana rows.",
            "Say each vowel evenly without stretching it too much.",
            "Read two-kana practice by saying one sound, then the next: あい = a-i.",
        ],
    },
    {
        "slug": "hira_02_k_row.docx",
        "title": "Hiragana Set 2 — K Row",
        "subtitle": "か き く け こ",
        "focus": "Keep the k sound the same and change only the vowel after it.",
        "kana": [
            ("か", "ka", "k + a"),
            ("き", "ki", "k + i"),
            ("く", "ku", "k + u"),
            ("け", "ke", "k + e"),
            ("こ", "ko", "k + o"),
        ],
        "pairs": [("かき", "ka-ki"), ("くけ", "ku-ke"), ("ここ", "ko-ko"), ("かこ", "ka-ko")],
        "teaching_notes": [
            "The consonant stays k all the way through the row.",
            "Only the vowel changes: a, i, u, e, o.",
            "Practice saying the row in order until it sounds smooth.",
        ],
    },
    {
        "slug": "hira_03_s_row.docx",
        "title": "Hiragana Set 3 — S Row",
        "subtitle": "さ し す せ そ",
        "focus": "The s row introduces し, which sounds like shi instead of si.",
        "kana": [
            ("さ", "sa", "s + a"),
            ("し", "shi", "Sounds like shee."),
            ("す", "su", "s + u"),
            ("せ", "se", "s + e"),
            ("そ", "so", "s + o"),
        ],
        "pairs": [("さし", "sa-shi"), ("すし", "su-shi"), ("せそ", "se-so"), ("そさ", "so-sa")],
        "teaching_notes": [
            "し is one of the first kana that does not match a simple consonant + vowel pattern.",
            "Students often remember すし because sushi is already familiar in English.",
            "Keep the vowel steady even when the spelling changes.",
        ],
    },
    {
        "slug": "hira_04_t_row.docx",
        "title": "Hiragana Set 4 — T Row",
        "subtitle": "た ち つ て と",
        "focus": "The t row has two special sounds: ち = chi and つ = tsu.",
        "kana": [
            ("た", "ta", "t + a"),
            ("ち", "chi", "Sounds like chee."),
            ("つ", "tsu", "Starts with ts sound."),
            ("て", "te", "t + e"),
            ("と", "to", "t + o"),
        ],
        "pairs": [("たち", "ta-chi"), ("つて", "tsu-te"), ("とた", "to-ta"), ("てと", "te-to")],
        "teaching_notes": [
            "ち and つ are common beginner trouble spots, so say them slowly first.",
            "つ begins with a small ts sound before the vowel u.",
            "Mixing this row with the s row is normal at first, so practice aloud often.",
        ],
    },
    {
        "slug": "hira_05_n_row.docx",
        "title": "Hiragana Set 5 — N Row",
        "subtitle": "な に ぬ ね の",
        "focus": "This row is regular and is good for building reading confidence.",
        "kana": [
            ("な", "na", "n + a"),
            ("に", "ni", "n + i"),
            ("ぬ", "nu", "n + u"),
            ("ね", "ne", "n + e"),
            ("の", "no", "n + o"),
        ],
        "pairs": [("なに", "na-ni"), ("ぬね", "nu-ne"), ("のに", "no-ni"), ("なの", "na-no")],
        "teaching_notes": [
            "This row follows the sound pattern very clearly, which makes it a good review row.",
            "に appears in many basic words, so it is useful to memorize early.",
            "Read each pair evenly and do not rush the second syllable.",
        ],
    },
    {
        "slug": "hira_06_h_row.docx",
        "title": "Hiragana Set 6 — H Row",
        "subtitle": "は ひ ふ へ ほ",
        "focus": "The h row is mostly regular, but ふ sounds like fu instead of hu.",
        "kana": [
            ("は", "ha", "h + a"),
            ("ひ", "hi", "h + i"),
            ("ふ", "fu", "Soft f sound with u."),
            ("へ", "he", "h + e"),
            ("ほ", "ho", "h + o"),
        ],
        "pairs": [("はひ", "ha-hi"), ("ふへ", "fu-he"), ("ほは", "ho-ha"), ("ひふ", "hi-fu")],
        "teaching_notes": [
            "ふ is pronounced more softly than an English full f sound.",
            "Later, は and へ can sound different in particles, but here use the regular row sound.",
            "This worksheet focuses only on the basic row reading.",
        ],
    },
    {
        "slug": "hira_07_m_row.docx",
        "title": "Hiragana Set 7 — M Row",
        "subtitle": "ま み む め も",
        "focus": "The m row is regular and easy to blend into simple syllable pairs.",
        "kana": [
            ("ま", "ma", "m + a"),
            ("み", "mi", "m + i"),
            ("む", "mu", "m + u"),
            ("め", "me", "m + e"),
            ("も", "mo", "m + o"),
        ],
        "pairs": [("まみ", "ma-mi"), ("むめ", "mu-me"), ("もま", "mo-ma"), ("みも", "mi-mo")],
        "teaching_notes": [
            "This is another strong confidence-building row because the sounds are very regular.",
            "Encourage students to read pairs in one smooth rhythm instead of stopping between kana.",
            "Quick oral drills work well with this row.",
        ],
    },
    {
        "slug": "hira_08_y_row.docx",
        "title": "Hiragana Set 8 — Y Row",
        "subtitle": "や ゆ よ",
        "focus": "The y row has only three kana, so students can focus on hearing the vowel clearly.",
        "kana": [
            ("や", "ya", "y + a"),
            ("ゆ", "yu", "y + u"),
            ("よ", "yo", "y + o"),
        ],
        "pairs": [("やゆ", "ya-yu"), ("ゆよ", "yu-yo"), ("よや", "yo-ya"), ("やよ", "ya-yo")],
        "teaching_notes": [
            "There is no yi or ye in the basic hiragana chart.",
            "Because the row is short, it works well as a quick review worksheet.",
            "Focus on keeping ya, yu, and yo clearly different from one another.",
        ],
    },
    {
        "slug": "hira_09_r_row.docx",
        "title": "Hiragana Set 9 — R Row",
        "subtitle": "ら り る れ ろ",
        "focus": "The Japanese r sound is light, somewhere between English r, l, and d.",
        "kana": [
            ("ら", "ra", "Light tapped r + a"),
            ("り", "ri", "Light tapped r + i"),
            ("る", "ru", "Light tapped r + u"),
            ("れ", "re", "Light tapped r + e"),
            ("ろ", "ro", "Light tapped r + o"),
        ],
        "pairs": [("らり", "ra-ri"), ("るれ", "ru-re"), ("ろら", "ro-ra"), ("りろ", "ri-ro")],
        "teaching_notes": [
            "The Japanese r is lighter than the English r sound.",
            "Students should try a quick tap with the tongue instead of a long hard r.",
            "Say the row slowly first, then build speed after the sound feels natural.",
        ],
    },
    {
        "slug": "hira_10_w_n.docx",
        "title": "Hiragana Set 10 — W Row and ん",
        "subtitle": "わ を ん",
        "focus": "This worksheet introduces the last basic kana in the beginner chart.",
        "kana": [
            ("わ", "wa", "w + a"),
            ("を", "o", "Usually pronounced o in modern Japanese."),
            ("ん", "n", "A stand-alone n sound."),
        ],
        "pairs": [("わを", "wa-o"), ("をん", "o-n"), ("わん", "wa-n"), ("んわ", "n-wa")],
        "teaching_notes": [
            "を is mainly used as a grammar particle, but beginners still learn it in the chart.",
            "ん is special because it is a sound on its own, not a consonant + vowel pair.",
            "This worksheet is a good review point for the full beginner hiragana chart.",
        ],
    },
]


def _rgb_hex(rgb):
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_bg(cell, rgb):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), _rgb_hex(rgb))
    tcPr.append(shd)


def set_para_bg(para, rgb):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), _rgb_hex(rgb))
    pPr.append(shd)


def add_run(para, text, bold=False, italic=False, size=12, color=DARK, font='Calibri'):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = font
    return r


def add_corner_watermark(doc, text='neuralforge.cc'):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        add_run(p, text, size=7, color=GREY)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, bold=True, size=14 if level == 1 else 12,
            color=NAVY if level == 1 else TEAL)
    return p


def section_banner(doc, text, bg=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    add_run(p, text, bold=True, size=12, color=WHITE)
    set_para_bg(p, bg)
    return p


def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    p.add_run()._r.append(br)


def add_page_label(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, size=10, color=GREY)


def add_info_row(doc):
    info = doc.add_table(rows=1, cols=4)
    info.style = 'Table Grid'
    for i, (lbl, blank) in enumerate([
        ('Name', '________________________'),
        ('Class', '___________'),
        ('Date', '______________'),
        ('Period', '____'),
    ]):
        c = info.rows[0].cells[i]
        set_cell_bg(c, LIGHT)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, f'{lbl}: ', bold=True, size=10, color=NAVY)
        add_run(p, blank, size=10, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def get_fill_in_prompts(config):
    kana = [item[0] for item in config['kana']]
    if len(kana) >= 5:
        return [
            (f"{kana[0]}  __  {kana[2]}  __  {kana[4]}", f"{kana[1]}, {kana[3]}"),
        ]
    return [
        (f"{kana[0]}  __  {kana[2]}", kana[1]),
    ]


def get_distractors(config):
    current = {kana for kana, _sound, _tip in config['kana']}
    idx = WORKSHEETS.index(config)
    candidates = []

    if idx > 0:
        candidates.extend(kana for kana, _sound, _tip in WORKSHEETS[idx - 1]['kana'])
    if idx + 1 < len(WORKSHEETS):
        candidates.extend(kana for kana, _sound, _tip in WORKSHEETS[idx + 1]['kana'])

    candidates.extend(
        kana
        for worksheet in WORKSHEETS
        for kana, _sound, _tip in worksheet['kana']
        if kana not in current
    )

    unique = []
    for kana in candidates:
        if kana not in current and kana not in unique:
            unique.append(kana)
    return unique[:3]


def get_odd_one_out_prompts(config):
    kana = [item[0] for item in config['kana']]
    distractors = get_distractors(config)
    if len(kana) >= 5:
        return [
            ([distractors[0], kana[0], kana[1], kana[2]], distractors[0]),
            ([kana[1], distractors[1], kana[2], kana[3]], distractors[1]),
        ]
    return [
        ([distractors[0], kana[0], kana[1], kana[2]], distractors[0]),
        ([kana[0], distractors[1], kana[1], kana[2]], distractors[1]),
    ]


def get_true_false_prompts(config):
    sounds = [sound for _kana, sound, _tip in config['kana']]
    if len(config['kana']) >= 5:
        return [
            (f"This set follows the sound order {', '.join(sounds)}.", 'True'),
            ('This set has only three kana.', 'False'),
        ]
    return [
        (f"This set has {len(config['kana'])} kana.", 'True'),
        ('This set uses all five vowel positions.', 'False'),
    ]


def get_example_source(config):
    current_idx = WORKSHEETS.index(config)
    source_idx = 2 if current_idx != 2 else 1
    return WORKSHEETS[source_idx]


def get_worked_examples(config):
    source = get_example_source(config)
    return [
        (f"{source['kana'][0][0]}  ->  {source['kana'][0][1]}", 'Example from a different row.'),
        (get_fill_in_prompts(source)[0][0], get_fill_in_prompts(source)[0][1]),
    ]


def get_row_pattern_note(config):
    sounds = [sound for _, sound, _tip in config['kana']]
    if config['title'].endswith('Vowels'):
        return "These five vowels are the sound base of the whole hiragana chart: a, i, u, e, o."
    if len(sounds) == 5:
        return f"Most full hiragana rows follow the vowel order a, i, u, e, o. In this set that sounds like: {', '.join(sounds)}."
    return f"This is a short row, so it does not use all five vowel spots. Memorize the order as: {', '.join(sounds)}."


def get_special_sound_note(config):
    sounds = [sound for _, sound, _tip in config['kana']]
    special_map = {
        'shi': 'し = shi',
        'chi': 'ち = chi',
        'tsu': 'つ = tsu',
        'fu': 'ふ = fu',
    }
    specials = [text for sound, text in special_map.items() if sound in sounds]
    if 'W Row and ん' in config['title']:
        return "This set has two special basics: を is usually read as o, and ん is its own sound by itself."
    if specials:
        return f"Watch the special reading in this set: {', '.join(specials)}."
    return "This set is very regular, so you can focus on the letter shape and the changing vowel sound."


def build_header(doc, config):
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = 'Table Grid'
    r0 = tbl.rows[0].cells[0]
    set_cell_bg(r0, NAVY)
    r0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = r0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(5)
    p0.paragraph_format.space_after = Pt(2)
    add_run(p0, 'Beginner Japanese', bold=True, size=20, color=WHITE)
    add_run(p0, f"\n{config['title']}", size=13, color=GOLD)

    r1 = tbl.rows[1].cells[0]
    set_cell_bg(r1, TEAL)
    p1 = r1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(3)
    p1.paragraph_format.space_after = Pt(3)
    add_run(p1, 'Student Worksheet  ·  Page 1 of 3', size=10, color=WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def build_page1(doc, config):
    build_header(doc, config)
    add_info_row(doc)

    heading(doc, 'How this set works')
    for tip in [
        f"Set: {config['subtitle']}. {config['focus']}",
        f"{get_row_pattern_note(config)} {get_special_sound_note(config)}",
        "Parts: write the reading, fill a pattern, find the odd one out, and answer true or false.",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        add_run(p, tip, size=8)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    section_banner(doc, 'Worked Examples', bg=GOLD)

    example_note = doc.add_paragraph()
    example_note.paragraph_format.space_before = Pt(0)
    example_note.paragraph_format.space_after = Pt(1)
    add_run(example_note, 'Format only example from another row: ', size=9, color=DARK)
    add_run(example_note, f"{get_example_source(config)['kana'][0][0]} = {get_example_source(config)['kana'][0][1]}", bold=True, size=9, color=TEAL)

    examples_tbl = doc.add_table(rows=1, cols=1)
    examples_tbl.style = 'Table Grid'
    source = get_example_source(config)
    prompt, answer = get_worked_examples(config)[1]
    cell = examples_tbl.rows[0].cells[0]
    set_cell_bg(cell, LIGHT)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, 'Pattern example', bold=True, size=8, color=TEAL)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, prompt, bold=True, size=9, color=NAVY)
    p3 = cell.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(1)
    add_run(p3, f'Answer: {answer}', size=8, color=DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    section_banner(doc, 'Unit 1 — Write the Reading', bg=TEAL)

    act1_intro = doc.add_paragraph()
    act1_intro.paragraph_format.space_before = Pt(0)
    act1_intro.paragraph_format.space_after = Pt(1)
    add_run(act1_intro, 'Write the romaji reading for each kana.', size=8, color=DARK)

    reading_tbl = doc.add_table(rows=len(config['kana']) + 1, cols=3)
    reading_tbl.style = 'Table Grid'
    for ci, label in enumerate(['#', 'Kana', 'Your reading']):
        c = reading_tbl.rows[0].cells[ci]
        set_cell_bg(c, NAVY)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        add_run(p, label, bold=True, size=8, color=WHITE)

    for idx, (kana, _sound, _tip) in enumerate(config['kana'], 1):
        row = reading_tbl.rows[idx]
        bg = LIGHT if idx % 2 == 1 else WHITE
        vals = [
            (str(idx), GREY, 9, True),
            (kana, NAVY, 14, True),
            ('__________________', DARK, 10, False),
        ]
        for ci, (val, clr, sz, bld) in enumerate(vals):
            c = row.cells[ci]
            set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, val, bold=bld, size=(8 if ci != 1 else 12), color=clr)


def build_page2(doc, config):
    page_break(doc)
    add_page_label(doc, 'Student Worksheet  ·  Page 2 of 3')

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    section_banner(doc, 'Unit 2 — Fill in the Missing Kana', bg=NAVY)

    act2_intro = doc.add_paragraph()
    act2_intro.paragraph_format.space_before = Pt(0)
    act2_intro.paragraph_format.space_after = Pt(1)
    add_run(act2_intro, 'Fill in the missing kana to complete the row.', size=8, color=DARK)

    fill_tbl = doc.add_table(rows=len(get_fill_in_prompts(config)) + 1, cols=2)
    fill_tbl.style = 'Table Grid'
    for ci, label in enumerate(['Fill in the missing kana', 'Your answer']):
        c = fill_tbl.rows[0].cells[ci]
        set_cell_bg(c, TEAL)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        add_run(p, label, bold=True, size=8, color=WHITE)

    for idx, (prompt, _answer) in enumerate(get_fill_in_prompts(config), 1):
        row = fill_tbl.rows[idx]
        bg = LIGHT if idx % 2 == 1 else WHITE
        for ci, val in enumerate([prompt, '__________________']):
            c = row.cells[ci]
            set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_run(p, val, size=8 if ci == 1 else 10, color=DARK if ci == 1 else NAVY, bold=(ci == 0))

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    section_banner(doc, 'Unit 3 — Find the Odd One Out', bg=TEAL)

    odd_intro = doc.add_paragraph()
    odd_intro.paragraph_format.space_before = Pt(0)
    odd_intro.paragraph_format.space_after = Pt(1)
    add_run(odd_intro, 'Write the one kana that does not belong.', size=8, color=DARK)

    odd_tbl = doc.add_table(rows=len(get_odd_one_out_prompts(config)) + 1, cols=2)
    odd_tbl.style = 'Table Grid'
    for ci, label in enumerate(['Circle the kana that does not belong', 'Your answer']):
        c = odd_tbl.rows[0].cells[ci]
        set_cell_bg(c, NAVY)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        add_run(p, label, bold=True, size=8, color=WHITE)

    for idx, (prompt_items, _answer) in enumerate(get_odd_one_out_prompts(config), 1):
        row = odd_tbl.rows[idx]
        bg = MINT if idx % 2 == 1 else WHITE
        for ci, val in enumerate(['   '.join(prompt_items), '________']):
            c = row.cells[ci]
            set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_run(p, val, size=8 if ci == 1 else 10, color=DARK if ci == 1 else NAVY, bold=(ci == 0))

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    section_banner(doc, 'Unit 4 — True or False', bg=GOLD)

    tf_intro = doc.add_paragraph()
    tf_intro.paragraph_format.space_before = Pt(0)
    tf_intro.paragraph_format.space_after = Pt(1)
    add_run(tf_intro, 'Write True or False for each sentence.', size=8, color=DARK)

    tf_tbl = doc.add_table(rows=len(get_true_false_prompts(config)) + 1, cols=2)
    tf_tbl.style = 'Table Grid'
    for ci, label in enumerate(['Statement', 'True / False']):
        c = tf_tbl.rows[0].cells[ci]
        set_cell_bg(c, NAVY)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        add_run(p, label, bold=True, size=8, color=WHITE)

    for idx, (statement, _answer) in enumerate(get_true_false_prompts(config), 1):
        row = tf_tbl.rows[idx]
        bg = LIGHT if idx % 2 == 1 else WHITE
        for ci, val in enumerate([statement, '________']):
            c = row.cells[ci]
            set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if ci == 0:
                add_run(p, val, size=7, color=DARK)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(p, val, size=8, color=DARK)


def build_page3(doc, config):
    page_break(doc)
    add_page_label(doc, 'Reference & Answer Key  ·  Page 3 of 3')
    section_banner(doc, 'Reference & Answer Key', bg=DARK)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(0)
    intro.paragraph_format.space_after = Pt(1)
    add_run(intro, 'Use this page as a reference after students finish page 1.', size=9, color=DARK)

    pair_tbl = doc.add_table(rows=len(config['pairs']) + 1, cols=2)
    pair_tbl.style = 'Table Grid'
    for ci, label in enumerate(['Practice pair', 'Reading']):
        c = pair_tbl.rows[0].cells[ci]
        set_cell_bg(c, NAVY)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        add_run(p, label, bold=True, size=9, color=WHITE)

    for idx, (pair, reading) in enumerate(config['pairs'][:2], 1):
        bg = LIGHT if idx % 2 == 1 else WHITE
        row = pair_tbl.rows[idx]
        for ci, (val, clr, sz, bld) in enumerate([
            (pair, NAVY, 12, True),
            (reading, DARK, 9, False),
        ]):
            c = row.cells[ci]
            set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_run(p, val, bold=bld, size=sz, color=clr)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    section_banner(doc, 'Unit 1 Answers', bg=TEAL)

    answer_tbl = doc.add_table(rows=len(config['kana']) + 1, cols=2)
    answer_tbl.style = 'Table Grid'
    for ci, label in enumerate(['Kana', 'Reading']):
        c = answer_tbl.rows[0].cells[ci]
        set_cell_bg(c, TEAL)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        add_run(p, label, bold=True, size=9, color=WHITE)

    for idx, (kana, sound, _tip) in enumerate(config['kana'], 1):
        bg = MINT if idx % 2 == 1 else WHITE
        row = answer_tbl.rows[idx]
        for ci, (val, clr, sz, bld) in enumerate([
            (kana, NAVY, 13, True),
            (sound, DARK, 10, False),
        ]):
            c = row.cells[ci]
            set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_run(p, val, bold=bld, size=(9 if ci == 1 else 12), color=clr)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    section_banner(doc, 'Unit 2 Answers', bg=NAVY)

    fill_answer_tbl = doc.add_table(rows=len(get_fill_in_prompts(config)) + 1, cols=2)
    fill_answer_tbl.style = 'Table Grid'
    for ci, label in enumerate(['Fill in the missing kana', 'Answer']):
        c = fill_answer_tbl.rows[0].cells[ci]
        set_cell_bg(c, NAVY)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        add_run(p, label, bold=True, size=9, color=WHITE)

    for idx, (prompt, answer) in enumerate(get_fill_in_prompts(config), 1):
        bg = LIGHT if idx % 2 == 1 else WHITE
        row = fill_answer_tbl.rows[idx]
        for ci, (val, clr, sz, bld) in enumerate([
            (prompt, NAVY, 11, True),
            (answer, DARK, 10, False),
        ]):
            c = row.cells[ci]
            set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_run(p, val, bold=bld, size=(9 if ci == 1 else 10), color=clr)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    section_banner(doc, 'Unit 3 Answers', bg=TEAL)

    odd_answer_tbl = doc.add_table(rows=len(get_odd_one_out_prompts(config)) + 1, cols=2)
    odd_answer_tbl.style = 'Table Grid'
    for ci, label in enumerate(['Circle the kana that does not belong', 'Answer']):
        c = odd_answer_tbl.rows[0].cells[ci]
        set_cell_bg(c, TEAL)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        add_run(p, label, bold=True, size=9, color=WHITE)

    for idx, (prompt_items, answer) in enumerate(get_odd_one_out_prompts(config), 1):
        bg = MINT if idx % 2 == 1 else WHITE
        row = odd_answer_tbl.rows[idx]
        for ci, (val, clr, sz, bld) in enumerate([
            ('   '.join(prompt_items), NAVY, 11, True),
            (answer, DARK, 10, False),
        ]):
            c = row.cells[ci]
            set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_run(p, val, bold=bld, size=(9 if ci == 1 else 10), color=clr)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    section_banner(doc, 'Unit 4 Answers', bg=GOLD)

    tf_answer_tbl = doc.add_table(rows=len(get_true_false_prompts(config)) + 1, cols=2)
    tf_answer_tbl.style = 'Table Grid'
    for ci, label in enumerate(['Statement', 'Answer']):
        c = tf_answer_tbl.rows[0].cells[ci]
        set_cell_bg(c, NAVY)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        add_run(p, label, bold=True, size=9, color=WHITE)

    for idx, (statement, answer) in enumerate(get_true_false_prompts(config), 1):
        row = tf_answer_tbl.rows[idx]
        bg = LIGHT if idx % 2 == 1 else WHITE
        for ci, val in enumerate([statement, answer]):
            c = row.cells[ci]
            set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if ci == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(p, val, bold=True, size=9, color=TEAL)
            else:
                add_run(p, val, size=8, color=DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    heading(doc, 'Quick review', level=2)
    quick_notes = [
        'One hiragana usually represents one sound beat.',
        'Most rows follow the vowel order a, i, u, e, o.',
        "Point to one letter at a time and say the full sound in one step.",
        config['teaching_notes'][0],
    ]
    for tip in quick_notes:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        add_run(p, '▸ ', bold=True, size=9, color=GOLD)
        add_run(p, tip, size=9, color=DARK)


def build_doc(config):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.2)
        section.bottom_margin = Inches(0.2)
        section.left_margin = Inches(0.2)
        section.right_margin = Inches(0.2)

    add_corner_watermark(doc)
    build_page1(doc, config)
    build_page2(doc, config)
    build_page3(doc, config)
    return doc


def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))

    for config in WORKSHEETS:
        doc = build_doc(config)
        docx_path, pdf_path = save_docx_and_pdf(doc, script_dir, config['slug'], subdir='hiragana_series')
        print(f'Saved DOCX: {docx_path}')
        print(f'Saved PDF: {pdf_path}')


if __name__ == '__main__':
    main()
