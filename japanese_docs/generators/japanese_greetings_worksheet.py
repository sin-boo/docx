import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from output_utils import save_outputs

# ── Colour palette ─────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1A, 0x3A, 0x5C)
TEAL  = RGBColor(0x00, 0x7A, 0x87)
GOLD  = RGBColor(0xE6, 0xA8, 0x17)
LIGHT = RGBColor(0xF0, 0xF6, 0xFA)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
GREY  = RGBColor(0x6B, 0x7B, 0x8D)

# ── Content data ───────────────────────────────────────────────────────────────
QUESTIONS = [
    ("Greeting someone in the morning",               "ohayou"),
    ("Greeting someone in the afternoon / daytime",   "konnichiwa"),
    ("Greeting someone in the evening / night",       "konbanwa"),
    ("Saying hello when meeting someone",             "konnichiwa"),
    ("Saying goodbye",                                "sayounara"),
    ("Asking someone 'Excuse me' before a question",  "sumimasen"),
    ("Asking someone for directions",                 "sumimasen, michi o oshiete kudasai"),
    ("Saying thank you",                              "arigatou"),
    ("Saying 'thank you very much'",                  "arigatou gozaimasu"),
    ("Saying sorry / excuse me",                      "gomen nasai"),
]

SCENARIOS = [
    ('Scenario A — Meeting & Asking Directions', [
        ('A', 'greeting'),
        ('B', 'response'),
        ('A', 'ask for directions'),
        ('B', 'short direction'),
    ]),
    ('Scenario B — Help & Thanks', [
        ('A', 'asks for help'),
        ('B', 'helps'),
        ('A', 'says thank you'),
        ('B', 'polite reply'),
    ]),
]

ANSWERS = [
    ("1. Morning greeting",        "おはよう",                          "ohayou",                             "Casual. Polite: おはようございます (ohayou gozaimasu)."),
    ("2. Afternoon greeting",      "こんにちは",                        "konnichiwa",                         "General daytime greeting."),
    ("3. Evening greeting",        "こんばんは",                        "konbanwa",                           "Used after sunset / in the evening."),
    ("4. Hello when meeting",      "こんにちは",                        "konnichiwa",                         "Same as #2 for daytime meetings."),
    ("5. Goodbye",                 "さようなら",                        "sayounara",                          "Formal; used when parting."),
    ("6. Excuse me (before Q)",    "すみません",                        "sumimasen",                          "Also used to get attention or lightly apologize."),
    ("7. Asking for directions",   "すみません、みちをおしえてください", "sumimasen, michi o oshiete kudasai", "Polite request: 'Please tell me the way.'"),
    ("8. Thank you",               "ありがとう",                        "arigatou",                           "Casual. Polite: ありがとうございます (arigatou gozaimasu)."),
    ("9. Thank you very much",     "ありがとうございます",              "arigatou gozaimasu",                 "Polite; appropriate in most situations."),
    ("10. Sorry / excuse me",      "ごめんなさい",                      "gomen nasai",                        "Used to apologize; casual to neutral tone."),
]

TEACHER_ACTIVITIES = [
    ('Pronunciation drill (5 min)',  'Teacher says each phrase; class repeats 3×.'),
    ('Pair roleplay (10–12 min)',    'Students practice Parts A/B and swap roles.'),
    ('Quick written quiz (5 min)',   'Teacher reads English prompts; students write Japanese.'),
    ('Homework extension',           'Students submit two polite phrases, e.g. はじめまして (hajimemashite).'),
]

PRINT_TIPS = [
    'Print double-sided: student page front, answer key back.',
    'Increase line spacing to 1.3–1.5 for extra handwriting space.',
    'Font: Calibri or Arial; Title 20–24 pt bold; Body 12 pt.',
]

# ── Low-level XML helpers ──────────────────────────────────────────────────────
def _rgb_hex(rgb):
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  _rgb_hex(rgb))
    tcPr.append(shd)

def set_para_bg(para, rgb):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  _rgb_hex(rgb))
    pPr.append(shd)

# ── Paragraph / run helpers ────────────────────────────────────────────────────
def add_run(para, text, bold=False, italic=False, size=12, color=DARK, font='Calibri'):
    r = para.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = font
    return r


def add_corner_watermark(doc, text='neuralforge.cc'):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        add_run(p, text, size=7, color=GREY)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6 if level == 1 else 4)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, text, bold=True, size=14 if level == 1 else 12,
            color=NAVY if level == 1 else TEAL)
    return p

def body(doc, text, indent=False, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    add_run(p, text, size=11)
    return p

def blank_line(doc, text, label, hint='', indent_cm=0.6, space_after=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(indent_cm)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    add_run(p, f"{label}  ", bold=True, size=11, color=TEAL)
    add_run(p, text, size=11)
    if hint:
        add_run(p, f"  ({hint})", italic=True, size=10, color=GREY)
    return p

def section_banner(doc, text, bg=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.3)
    add_run(p, text, bold=True, size=12, color=WHITE)
    set_para_bg(p, bg)
    return p

def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    p.add_run()._r.append(br)

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '007A87')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── Page builders ──────────────────────────────────────────────────────────────
def build_page1(doc):
    """Banner, student info, objectives, instructions, Part 1 questions."""
    # Header banner
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = 'Table Grid'
    r0 = tbl.rows[0].cells[0]
    set_cell_bg(r0, NAVY)
    r0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = r0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before, p0.paragraph_format.space_after = Pt(5), Pt(2)
    add_run(p0, 'Beginner Japanese', bold=True, size=20, color=WHITE)
    add_run(p0, '\nGreetings & Essential Phrases', bold=False, size=13, color=GOLD)
    r1 = tbl.rows[1].cells[0]
    set_cell_bg(r1, TEAL)
    p1 = r1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before, p1.paragraph_format.space_after = Pt(3), Pt(3)
    add_run(p1, 'Student Worksheet  ·  Page 1 of 3', bold=False, size=10, color=WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Student info row
    info = doc.add_table(rows=1, cols=4)
    info.style = 'Table Grid'
    for i, (lbl, blank) in enumerate([
        ('Name',   '_________________________________'),
        ('Class',  '___________'),
        ('Date',   '__________________'),
        ('Period', '____'),
    ]):
        c = info.rows[0].cells[i]
        set_cell_bg(c, LIGHT)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(2)
        add_run(p, f'{lbl}: ', bold=True, size=10, color=NAVY)
        add_run(p, blank,      bold=False, size=10, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Learning objectives
    heading(doc, 'Learning Objectives')
    for obj in [
        'Recognize and write basic Japanese greetings (hiragana).',
        'Practice pronunciation using romaji hints.',
        'Use phrases in short roleplay dialogues.',
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent  = Cm(0.6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        add_run(p, obj, size=10)

    # Instructions
    heading(doc, 'Instructions')
    body(doc, 'For each English prompt below:')
    for n, txt in [
        ('1.', 'Write the Japanese (hiragana/kanji) on the first line.'),
        ('2.', 'Write the romaji on the second line.  Romaji hints are in parentheses.'),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        add_run(p, f'{n} ', bold=True, size=10, color=TEAL)
        add_run(p, txt,     size=10)
    divider(doc)

    # Part 1 questions
    section_banner(doc, 'Part 1 — Greetings & Phrases', bg=NAVY)
    for i, (prompt, hint) in enumerate(QUESTIONS, 1):
        pq = doc.add_paragraph()
        pq.paragraph_format.space_before = Pt(2)
        pq.paragraph_format.space_after  = Pt(0)
        add_run(pq, f'{i}.  ', bold=True, size=10, color=TEAL)
        add_run(pq, prompt,    bold=False, size=10, color=DARK)
        blank_line(doc, '_' * 46, 'Japanese:', hint='',   indent_cm=0.8, space_after=0)
        blank_line(doc, '_' * 32, 'Romaji:  ', hint=hint, indent_cm=0.8, space_after=0)


def build_page2(doc):
    """Part 2 roleplay scenarios and pronunciation tips."""
    page_break(doc)

    # Page header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(0), Pt(2)
    add_run(p, 'Student Worksheet  ·  Page 2 of 3', bold=False, size=10, color=GREY)

    # Roleplay scenarios
    section_banner(doc, 'Part 2 — Short Roleplay Practice  (pair up)', bg=TEAL)
    for title, lines in SCENARIOS:
        ps = doc.add_paragraph()
        ps.paragraph_format.space_before = Pt(4)
        ps.paragraph_format.space_after  = Pt(1)
        add_run(ps, title, bold=True, size=10, color=NAVY)
        for speaker, cue in lines:
            pl = doc.add_paragraph()
            pl.paragraph_format.left_indent  = Cm(0.8)
            pl.paragraph_format.space_before = Pt(1)
            pl.paragraph_format.space_after  = Pt(1)
            add_run(pl, f'{speaker}:  ', bold=True, size=10, color=TEAL)
            add_run(pl, '_' * 38 + '  ', size=10)
            add_run(pl, f'({cue})',       italic=True, size=9, color=GREY)
    divider(doc)

    # Pronunciation tips
    heading(doc, 'Pronunciation Tips')
    for tip in [
        'Vowels: a / e / i / o / u — short and steady.',
        'ん (n) is nasal — hold it slightly.',
        'Repeat slowly 3×, then 2× at normal speed.',
    ]:
        pt = doc.add_paragraph()
        pt.paragraph_format.left_indent  = Cm(0.6)
        pt.paragraph_format.space_before = Pt(0)
        pt.paragraph_format.space_after  = Pt(1)
        add_run(pt, '▸  ', bold=True, size=10, color=GOLD)
        add_run(pt, tip,   size=10)


def build_page3(doc):
    """Answer key and teacher notes (teacher copy)."""
    page_break(doc)

    # Answer key banner
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = 'Table Grid'
    ra0 = tbl.rows[0].cells[0]
    set_cell_bg(ra0, DARK)
    ra0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    pa0 = ra0.paragraphs[0]
    pa0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pa0.paragraph_format.space_before, pa0.paragraph_format.space_after = Pt(5), Pt(2)
    add_run(pa0, 'Answer Key — Teacher Copy', bold=True, size=16, color=WHITE)
    add_run(pa0, '\nPage 3 of 3',             bold=False, size=11, color=GOLD)
    ra1 = tbl.rows[1].cells[0]
    set_cell_bg(ra1, GOLD)
    pa1 = ra1.paragraphs[0]
    pa1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pa1.paragraph_format.space_before, pa1.paragraph_format.space_after = Pt(2), Pt(2)
    add_run(pa1, 'For teacher use only — do not distribute to students', bold=False, size=9, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Answer table
    atbl = doc.add_table(rows=1, cols=4)
    atbl.style = 'Table Grid'
    for i, col_label in enumerate(['Prompt', 'Japanese', 'Romaji', 'Notes']):
        c = atbl.rows[0].cells[i]
        set_cell_bg(c, NAVY)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(2)
        add_run(p, col_label, bold=True, size=10, color=WHITE)
    for row_i, (prompt, jp, ro, note) in enumerate(ANSWERS):
        row = atbl.add_row()
        bg = LIGHT if row_i % 2 == 0 else WHITE
        for ci, (val, clr, sz, bld) in enumerate([
            (prompt, DARK, 10, False),
            (jp,     TEAL, 10, True),
            (ro,     DARK, 10, False),
            (note,   GREY,  9, False),
        ]):
            c = row.cells[ci]
            set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(2)
            add_run(p, val, bold=bld, size=sz, color=clr)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    divider(doc)

    # Teacher activities
    heading(doc, "Teacher's Notes & Suggested Activities")
    for title, desc in TEACHER_ACTIVITIES:
        pa = doc.add_paragraph()
        pa.paragraph_format.left_indent  = Cm(0.6)
        pa.paragraph_format.space_before = Pt(1)
        pa.paragraph_format.space_after  = Pt(1)
        add_run(pa, f'{title}: ', bold=True,  size=10, color=NAVY)
        add_run(pa, desc,         bold=False, size=10, color=DARK)
    divider(doc)

    # Print tips
    heading(doc, 'Printing & Formatting Tips')
    for t in PRINT_TIPS:
        pf = doc.add_paragraph()
        pf.paragraph_format.left_indent  = Cm(0.6)
        pf.paragraph_format.space_before = Pt(0)
        pf.paragraph_format.space_after  = Pt(1)
        add_run(pf, '▸  ', bold=True, size=10, color=GOLD)
        add_run(pf, t,     size=10)


# ── Entry point ────────────────────────────────────────────────────────────────
def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    add_corner_watermark(doc)

    build_page1(doc)
    build_page2(doc)
    build_page3(doc)

    script_dir = os.path.dirname(os.path.abspath(__file__))
    docx_path, pdf_path = save_outputs(doc, script_dir, 'japanese_greetings_worksheet.docx')
    print(f'Saved DOCX: {docx_path}')
    print(f'Saved PDF: {pdf_path}')


if __name__ == '__main__':
    main()
