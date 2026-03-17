import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from output_utils import save_outputs

# ── Colour palette ─────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1A, 0x3A, 0x5C)
TEAL  = RGBColor(0x00, 0x7A, 0x87)
GOLD  = RGBColor(0xE6, 0xA8, 0x17)
LIGHT = RGBColor(0xF0, 0xF6, 0xFA)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
GREY  = RGBColor(0x6B, 0x7B, 0x8D)
MINT  = RGBColor(0xD4, 0xF1, 0xF4)

# ── Content ────────────────────────────────────────────────────────────────────
PHRASE_BANK = [
    ("こんにちは", "Hello / Good afternoon"),
    ("はじめまして", "Nice to meet you"),
    ("げんきです", "I am fine / well"),
    ("あなたは？", "And you?"),
    ("はい、どうぞ", "Yes, here you go"),
    ("ありがとう", "Thank you"),
    ("どういたしまして", "You're welcome"),
    ("また あした", "See you tomorrow"),
]

DIALOGUES = [
    {
        "title": "1. Meeting someone new",
        "situation": "You meet a new classmate for the first time.",
        "a_en": "Nice to meet you. I'm Alex.",
        "b_jp": "はじめまして。わたしは ゆき です。",
        "meaning_answer": "Nice to meet you. I am Yuki.",
        "reply_answer": "はじめまして。わたしは アレックス です。",
        "note": "はじめまして is a common first-meeting phrase. わたしは ... です means 'I am ...'.",
    },
    {
        "title": "2. A simple greeting",
        "situation": "You see your friend at school.",
        "a_en": "Hi, Mika. How are you?",
        "b_jp": "こんにちは。げんきです。あなたは？",
        "meaning_answer": "Hello. I am fine. And you?",
        "reply_answer": "わたしも げんきです。",
        "note": "げんきです means 'I am fine / well.' あなたは？ means 'And you?'.",
    },
    {
        "title": "3. Borrowing a school item",
        "situation": "You ask to borrow an eraser.",
        "a_en": "Can I borrow an eraser?",
        "b_jp": "はい、どうぞ。",
        "meaning_answer": "Yes, here you go.",
        "reply_answer": "ありがとう。",
        "note": "どうぞ is useful when you give something or invite someone to go ahead.",
    },
    {
        "title": "4. Saying goodbye",
        "situation": "School is over and you are leaving.",
        "a_en": "I have to go now. Bye!",
        "b_jp": "また あした。",
        "meaning_answer": "See you tomorrow.",
        "reply_answer": "また あした。",
        "note": "また means 'again' and あした means 'tomorrow'.",
    },
]

BEGINNER_TIPS = [
    "Read the English line first so you know the situation.",
    "Then read the Japanese line slowly and look for phrases from the help box.",
    "After that, write the meaning in English and try a short Japanese reply.",
    "Short answers are okay for beginners. A correct simple reply is better than a long confusing one.",
]


# ── Helpers ────────────────────────────────────────────────────────────────────
def _rgb_hex(rgb):
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_bg(cell, rgb):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), _rgb_hex(rgb))
    tcPr.append(shd)


def set_para_bg(para, rgb):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), _rgb_hex(rgb))
    pPr.append(shd)


def remove_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right'):
        border = tcBorders.find(qn(f'w:{edge}'))
        if border is None:
            border = OxmlElement(f'w:{edge}')
            tcBorders.append(border)
        border.set(qn('w:val'), 'nil')


def add_run(para, text, bold=False, italic=False, size=12, color=DARK, font='Calibri'):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = font
    return r


def add_corner_watermark(doc, text='neuralforge.cc'):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        add_run(p, text, size=7, color=GREY)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, bold=True, size=14 if level == 1 else 12,
            color=NAVY if level == 1 else TEAL)
    return p


def section_banner(doc, text, bg=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    add_run(p, text, bold=True, size=12, color=WHITE)
    set_para_bg(p, bg)
    return p


def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    p.add_run()._r.append(br)


def add_page_label(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, size=10, color=GREY)


# ── Layout builders ────────────────────────────────────────────────────────────
def build_page1(doc):
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = 'Table Grid'
    r0 = tbl.rows[0].cells[0]
    set_cell_bg(r0, NAVY)
    r0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = r0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(5)
    p0.paragraph_format.space_after = Pt(2)
    add_run(p0, 'Beginner Japanese', bold=True, size=20, color=WHITE)
    add_run(p0, '\nSimple Dialogues & Replies', size=13, color=GOLD)

    r1 = tbl.rows[1].cells[0]
    set_cell_bg(r1, TEAL)
    p1 = r1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(3)
    p1.paragraph_format.space_after = Pt(3)
    add_run(p1, 'Student Worksheet  ·  Page 1 of 2', size=10, color=WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    info = doc.add_table(rows=1, cols=4)
    info.style = 'Table Grid'
    for i, (lbl, blank) in enumerate([
        ('Name', '________________________'),
        ('Class', '___________'),
        ('Date', '______________'),
        ('Period', '____'),
    ]):
        c = info.rows[0].cells[i]
        set_cell_bg(c, LIGHT)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, f'{lbl}: ', bold=True, size=10, color=NAVY)
        add_run(p, blank, size=10, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    heading(doc, 'How to use this worksheet')
    for tip in BEGINNER_TIPS:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        add_run(p, tip, size=10)

    section_banner(doc, 'Helpful Phrases', bg=TEAL)
    phrase_tbl = doc.add_table(rows=4, cols=2)
    phrase_tbl.style = 'Table Grid'
    phrase_tbl.autofit = False
    for row in phrase_tbl.rows:
        for cell in row.cells:
            cell.width = Inches(4.0)

    for idx, (jp, meaning) in enumerate(PHRASE_BANK):
        row_idx = idx // 2
        col_idx = idx % 2
        cell = phrase_tbl.rows[row_idx].cells[col_idx]
        set_cell_bg(cell, LIGHT if row_idx % 2 == 0 else WHITE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, jp, bold=True, size=11, color=NAVY)
        add_run(p, f'\n{meaning}', size=9, color=DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    section_banner(doc, 'Dialogue Practice', bg=NAVY)

    cards = doc.add_table(rows=2, cols=2)
    cards.style = 'Table Grid'
    cards.autofit = False

    for row in cards.rows:
        for cell in row.cells:
            cell.width = Inches(4.0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    for idx, card in enumerate(DIALOGUES):
        row_idx = idx // 2
        col_idx = idx % 2
        cell = cards.rows[row_idx].cells[col_idx]
        set_cell_bg(cell, MINT if row_idx % 2 == 0 else WHITE)

        p_title = cell.paragraphs[0]
        p_title.paragraph_format.space_before = Pt(2)
        p_title.paragraph_format.space_after = Pt(0)
        add_run(p_title, card['title'], bold=True, size=10, color=NAVY)

        p_situation = cell.add_paragraph()
        p_situation.paragraph_format.space_before = Pt(0)
        p_situation.paragraph_format.space_after = Pt(1)
        add_run(p_situation, 'Situation: ', bold=True, size=9, color=TEAL)
        add_run(p_situation, card['situation'], size=9, color=DARK)

        p_a = cell.add_paragraph()
        p_a.paragraph_format.space_before = Pt(0)
        p_a.paragraph_format.space_after = Pt(1)
        add_run(p_a, 'A (English): ', bold=True, size=9, color=TEAL)
        add_run(p_a, card['a_en'], size=9, color=DARK)

        p_b = cell.add_paragraph()
        p_b.paragraph_format.space_before = Pt(0)
        p_b.paragraph_format.space_after = Pt(1)
        add_run(p_b, 'B (Japanese): ', bold=True, size=9, color=TEAL)
        add_run(p_b, card['b_jp'], bold=True, size=10, color=NAVY)

        p_meaning = cell.add_paragraph()
        p_meaning.paragraph_format.space_before = Pt(0)
        p_meaning.paragraph_format.space_after = Pt(0)
        add_run(p_meaning, 'What did B say in English?', bold=True, size=8, color=DARK)

        p_meaning_line = cell.add_paragraph()
        p_meaning_line.paragraph_format.space_before = Pt(0)
        p_meaning_line.paragraph_format.space_after = Pt(1)
        add_run(p_meaning_line, '_' * 36, size=8, color=DARK)

        p_reply = cell.add_paragraph()
        p_reply.paragraph_format.space_before = Pt(0)
        p_reply.paragraph_format.space_after = Pt(0)
        add_run(p_reply, 'Your short Japanese reply:', bold=True, size=8, color=DARK)

        p_reply_line = cell.add_paragraph()
        p_reply_line.paragraph_format.space_before = Pt(0)
        p_reply_line.paragraph_format.space_after = Pt(2)
        add_run(p_reply_line, '_' * 28, size=8, color=DARK)


def build_page2(doc):
    page_break(doc)
    add_page_label(doc, 'Answer Key  ·  Page 2 of 2')
    section_banner(doc, 'Dialogue Answers & Beginner Notes', bg=DARK)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(1)
    intro.paragraph_format.space_after = Pt(2)
    add_run(intro, 'Each dialogue has one possible reply. Students may write other simple correct replies too.', size=10, color=DARK)

    ans_tbl = doc.add_table(rows=1, cols=4)
    ans_tbl.style = 'Table Grid'
    for i, col_label in enumerate(['Dialogue', 'Japanese line', 'Meaning in English', 'Sample reply']):
        c = ans_tbl.rows[0].cells[i]
        set_cell_bg(c, NAVY)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, col_label, bold=True, size=10, color=WHITE)

    for idx, card in enumerate(DIALOGUES):
        row = ans_tbl.add_row()
        bg = LIGHT if idx % 2 == 0 else WHITE
        values = [
            card['title'],
            card['b_jp'],
            card['meaning_answer'],
            card['reply_answer'],
        ]
        styles = [
            (DARK, 9, False),
            (TEAL, 10, True),
            (DARK, 9, False),
            (NAVY, 9, False),
        ]
        for ci, val in enumerate(values):
            c = row.cells[ci]
            set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            color, size, bold = styles[ci]
            add_run(p, val, bold=bold, size=size, color=color)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    heading(doc, 'Why these replies work')
    for idx, card in enumerate(DIALOGUES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        add_run(p, f'{idx}. ', bold=True, size=10, color=TEAL)
        add_run(p, card['note'], size=10, color=DARK)

    heading(doc, 'Teaching / Practice Ideas', level=2)
    for tip in [
        'Read each B line aloud twice and let students repeat it together.',
        'Let students compare English meanings with a partner before checking the answer key.',
        'Have students act out the dialogues and then replace one word, such as a name or school item.',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        add_run(p, '▸ ', bold=True, size=10, color=GOLD)
        add_run(p, tip, size=10, color=DARK)


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.2)
        section.bottom_margin = Inches(0.2)
        section.left_margin = Inches(0.2)
        section.right_margin = Inches(0.2)

    add_corner_watermark(doc)

    build_page1(doc)
    build_page2(doc)

    script_dir = os.path.dirname(os.path.abspath(__file__))
    docx_path, pdf_path = save_outputs(doc, script_dir, 'japanese_dialogue_worksheet.docx')
    print(f'Saved DOCX: {docx_path}')
    print(f'Saved PDF: {pdf_path}')


if __name__ == '__main__':
    main()
