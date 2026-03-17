import os
import random
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette (matches greetings worksheet) ───────────────────────────────
NAVY  = RGBColor(0x1A, 0x3A, 0x5C)
TEAL  = RGBColor(0x00, 0x7A, 0x87)
GOLD  = RGBColor(0xE6, 0xA8, 0x17)
LIGHT = RGBColor(0xF0, 0xF6, 0xFA)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
GREY  = RGBColor(0x6B, 0x7B, 0x8D)
MINT  = RGBColor(0xD4, 0xF1, 0xF4)   # soft teal tint for alternating rows

# ── Number data ────────────────────────────────────────────────────────────────
# (english_label, kanji, hiragana, romaji, notes)
NUMBERS = [
    ("1  — One",        "一",     "いち",         "ichi",        ""),
    ("2  — Two",        "二",     "に",           "ni",          ""),
    ("3  — Three",      "三",     "さん",         "san",         ""),
    ("4  — Four",       "四",     "し / よん",    "shi / yon",   "Both are used; よん is safer in daily speech."),
    ("5  — Five",       "五",     "ご",           "go",          ""),
    ("6  — Six",        "六",     "ろく",         "roku",        ""),
    ("7  — Seven",      "七",     "しち / なな",  "shichi / nana","なな is more common in daily speech."),
    ("8  — Eight",      "八",     "はち",         "hachi",       ""),
    ("9  — Nine",       "九",     "く / きゅう",  "ku / kyuu",   "きゅう is preferred in most contexts."),
    ("10 — Ten",        "十",     "じゅう",       "juu",         ""),
    ("11 — Eleven",     "十一",   "じゅういち",   "juu-ichi",    "Ten + one."),
    ("12 — Twelve",     "十二",   "じゅうに",     "juu-ni",      "Ten + two."),
    ("20 — Twenty",     "二十",   "にじゅう",     "ni-juu",      "Two × ten."),
    ("30 — Thirty",     "三十",   "さんじゅう",   "san-juu",     "Three × ten."),
    ("100 — One hundred","百",    "ひゃく",       "hyaku",       "New counter word."),
    ("1000 — One thousand","千",  "せん",         "sen",         "New counter word."),
]

# Matching page: focus on 1–20 in a compact two-column layout
MATCH_NUMBERS = [
    *NUMBERS[:12],  # 1–12
    ("13 — Thirteen", "十三", "じゅうさん",   "juu-san",   ""),
    ("14 — Fourteen", "十四", "じゅうよん",   "juu-yon",   ""),
    ("15 — Fifteen",  "十五", "じゅうご",     "juu-go",    ""),
    ("16 — Sixteen",  "十六", "じゅうろく",   "juu-roku",  ""),
    ("17 — Seventeen","十七", "じゅうなな",   "juu-nana",  ""),
    ("18 — Eighteen", "十八", "じゅうはち",   "juu-hachi", ""),
    ("19 — Nineteen", "十九", "じゅうきゅう", "juu-kyuu",  ""),
    NUMBERS[12],     # 20
]


# ── XML / formatting helpers (same as greetings file) ─────────────────────────
def _rgb_hex(rgb):
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

def set_cell_bg(cell, rgb):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  _rgb_hex(rgb))
    tcPr.append(shd)

def set_para_bg(para, rgb):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  _rgb_hex(rgb))
    pPr.append(shd)

def add_run(para, text, bold=False, italic=False, size=12, color=DARK, font='Calibri'):
    r = para.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size  = Pt(size)
    r.font.color.rgb = color
    r.font.name  = font
    return r

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6 if level == 1 else 4)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, text, bold=True,
            size=14 if level == 1 else 12,
            color=NAVY if level == 1 else TEAL)
    return p

def section_banner(doc, text, bg=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.3)
    add_run(p, text, bold=True, size=12, color=WHITE)
    set_para_bg(p, bg)
    return p

def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    p.add_run()._r.append(br)

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '007A87')
    pBdr.append(bot)
    pPr.append(pBdr)


def remove_cell_borders(cell):
    """Hide table cell borders for cleaner worksheet layouts."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right'):
        border = tcBorders.find(qn(f'w:{edge}'))
        if border is None:
            border = OxmlElement(f'w:{edge}')
            tcBorders.append(border)
        border.set(qn('w:val'), 'nil')
def get_matching_bank():
    """Return a deterministic answer bank and lookup for Part 2."""
    items = list(MATCH_NUMBERS)
    rng = random.Random(42)
    rng.shuffle(items)

    bank = []
    answer_lookup = {}
    for idx, item in enumerate(items):
        label = chr(ord('A') + idx)
        bank.append((label, item))
        answer_lookup[item[0]] = label
    return bank, answer_lookup


# ── Page builders ──────────────────────────────────────────────────────────────
TOTAL_PAGES = 7


def add_page_label(doc, page_no, total_pages, label='Student Worksheet'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, f'{label}  ·  Page {page_no} of {total_pages}', size=10, color=GREY)


def add_student_info(doc):
    info = doc.add_table(rows=1, cols=4)
    info.style = 'Table Grid'
    for i, (lbl, blank) in enumerate([
        ('Name',   '_________________________________'),
        ('Class',  '___________'),
        ('Date',   '__________________'),
        ('Period', '____'),
    ]):
        c = info.rows[0].cells[i]
        set_cell_bg(c, LIGHT)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(2)
        add_run(p, f'{lbl}: ', bold=True, size=10, color=NAVY)
        add_run(p, blank, size=10, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_title_banner(doc, total_pages):
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = 'Table Grid'
    r0 = tbl.rows[0].cells[0]
    set_cell_bg(r0, NAVY)
    r0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = r0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before, p0.paragraph_format.space_after = Pt(5), Pt(2)
    add_run(p0, 'Beginner Japanese', bold=True, size=20, color=WHITE)
    add_run(p0, '\nNumbers 1–1000  ·  数字 (sūji)', size=13, color=GOLD)

    r1 = tbl.rows[1].cells[0]
    set_cell_bg(r1, TEAL)
    p1 = r1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before, p1.paragraph_format.space_after = Pt(3), Pt(3)
    add_run(p1, f'Student Worksheet  ·  Page 1 of {total_pages}', size=10, color=WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_writing_page(doc, page_no, total_pages, title, answer_label):
    if page_no == 1:
        add_title_banner(doc, total_pages)
        add_student_info(doc)
    else:
        page_break(doc)
        add_page_label(doc, page_no, total_pages)

    section_banner(doc, title, bg=NAVY if answer_label == 'Kanji' else TEAL)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, 'Write only ', size=10)
    add_run(p, answer_label, bold=True, size=10, color=TEAL)
    add_run(p, ' for each English number below.', size=10)

    grid = doc.add_table(rows=(len(NUMBERS) + 1) // 2, cols=2)
    grid.style = 'Table Grid'
    grid.autofit = False
    col_width = Cm(8.3)

    for row_idx, row in enumerate(grid.rows):
        for col_idx, cell in enumerate(row.cells):
            item_idx = row_idx * 2 + col_idx
            remove_cell_borders(cell)
            cell.width = col_width
            set_cell_bg(cell, LIGHT if row_idx % 2 == 0 else WHITE)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            if item_idx >= len(NUMBERS):
                cell.text = ''
                continue

            eng, _kanji, _hira, _ro, _note = NUMBERS[item_idx]

            p_q = cell.paragraphs[0]
            p_q.paragraph_format.space_before = Pt(2)
            p_q.paragraph_format.space_after = Pt(0)
            add_run(p_q, f'{item_idx + 1}. ', bold=True, size=10, color=TEAL)
            add_run(p_q, eng, size=10, color=DARK)

            p_line = cell.add_paragraph()
            p_line.paragraph_format.left_indent = Cm(0.2)
            p_line.paragraph_format.space_before = Pt(0)
            p_line.paragraph_format.space_after = Pt(1)
            add_run(p_line, f'{answer_label}: ', bold=True, size=9, color=TEAL)
            add_run(p_line, '_' * 22, size=9, color=DARK)


def build_page1(doc):
    build_writing_page(doc, 1, TOTAL_PAGES, 'Page 1 — Kanji Writing Practice', 'Kanji')


def build_page2(doc):
    build_writing_page(doc, 2, TOTAL_PAGES, 'Page 2 — Hiragana Writing Practice', 'Hiragana')


def build_page3(doc):
    page_break(doc)
    add_page_label(doc, 3, TOTAL_PAGES)

    section_banner(doc, 'Page 3 — Match the Number Meanings', bg=TEAL)

    for n, txt in [
        ('1.', 'Look at the kanji and hiragana in each box.'),
        ('2.', 'Find the matching English word in the answer bank.'),
        ('3.', 'Write the correct letter on the blank line.'),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        add_run(p, f'{n} ', bold=True, size=10, color=TEAL)
        add_run(p, txt, size=10)

    bank_intro = doc.add_paragraph()
    bank_intro.paragraph_format.space_before = Pt(1)
    bank_intro.paragraph_format.space_after = Pt(2)
    add_run(bank_intro, 'Answer bank', bold=True, size=11, color=NAVY)
    add_run(bank_intro, '  Use the letters below.', size=10, color=DARK)

    bank, _answer_lookup = get_matching_bank()
    bank_tbl = doc.add_table(rows=5, cols=4)
    bank_tbl.style = 'Table Grid'
    for idx, (label, item) in enumerate(bank):
        row_idx = idx // 4
        col_idx = idx % 4
        eng_display = item[0].split('—')[-1].strip()
        cell = bank_tbl.rows[row_idx].cells[col_idx]
        set_cell_bg(cell, LIGHT if row_idx % 2 == 0 else WHITE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(2)
        add_run(p, f'{label}. ', bold=True, size=10, color=TEAL)
        add_run(p, eng_display, size=9, color=DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    grid_tbl = doc.add_table(rows=10, cols=2)
    grid_tbl.style = 'Table Grid'
    for pair_idx in range(10):
        for col_idx in range(2):
            item_idx = pair_idx * 2 + col_idx
            eng, kanji, hira, _ro, _note = MATCH_NUMBERS[item_idx]
            cell = grid_tbl.rows[pair_idx].cells[col_idx]
            set_cell_bg(cell, MINT if pair_idx % 2 == 0 else WHITE)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p_top = cell.paragraphs[0]
            p_top.paragraph_format.space_before = Pt(2)
            p_top.paragraph_format.space_after = Pt(0)
            add_run(p_top, f'{item_idx + 1}. ', bold=True, size=10, color=NAVY)
            add_run(p_top, 'Letter: ', bold=True, size=9, color=TEAL)
            add_run(p_top, '____', size=9, color=DARK)

            p_body = cell.add_paragraph()
            p_body.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_body.paragraph_format.space_before = Pt(1)
            p_body.paragraph_format.space_after = Pt(2)
            add_run(p_body, kanji, bold=True, size=18, color=NAVY)
            add_run(p_body, f'\n{hira}', size=10, color=TEAL)


def build_page4(doc):
    page_break(doc)
    add_page_label(doc, 4, TOTAL_PAGES)

    section_banner(doc, 'Page 4 — How Japanese Numbers Work', bg=NAVY)

    for tip in [
        'Base numbers: learn 1-10 first, then build larger numbers from those pieces.',
        '11-19 use 十 + the ones digit: 十一, 十二, 十三 ... 十九.',
        '20 and beyond follow the same pattern: 二十, 二十一, 三十五, 九十九, 百, 千.',
    ]:
        pt = doc.add_paragraph()
        pt.paragraph_format.left_indent = Cm(0.6)
        pt.paragraph_format.space_before = Pt(0)
        pt.paragraph_format.space_after = Pt(1)
        add_run(pt, '- ', bold=True, size=10, color=GOLD)
        add_run(pt, tip, size=9)

    ref = doc.add_paragraph()
    ref.paragraph_format.space_before = Pt(2)
    ref.paragraph_format.space_after = Pt(2)
    add_run(ref, 'Reference sheet: 1-20', bold=True, size=11, color=TEAL)

    ref_tbl = doc.add_table(rows=10, cols=2)
    ref_tbl.style = 'Table Grid'
    for pair_idx in range(10):
        for col_idx in range(2):
            item_idx = pair_idx * 2 + col_idx
            eng, kanji, hira, _ro, note = MATCH_NUMBERS[item_idx]
            cell = ref_tbl.rows[pair_idx].cells[col_idx]
            set_cell_bg(cell, LIGHT if pair_idx % 2 == 0 else WHITE)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p_eng = cell.paragraphs[0]
            p_eng.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_eng.paragraph_format.space_before = Pt(2)
            p_eng.paragraph_format.space_after = Pt(0)
            add_run(p_eng, eng.split('—')[-1].strip(), bold=True, size=10, color=DARK)

            p_body = cell.add_paragraph()
            p_body.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_body.paragraph_format.space_before = Pt(0)
            p_body.paragraph_format.space_after = Pt(2)
            add_run(p_body, kanji, bold=True, size=16, color=NAVY)
            add_run(p_body, f'\n{hira}', size=10, color=TEAL)
            if note:
                add_run(p_body, f'\n{note}', italic=True, size=7, color=GREY)


def build_page5(doc):
    page_break(doc)
    add_page_label(doc, 5, TOTAL_PAGES, 'Answer Key — Teacher Copy')

    section_banner(doc, 'Page 5 — Writing Answer Key', bg=DARK)

    atbl = doc.add_table(rows=1, cols=4)
    atbl.style = 'Table Grid'
    for i, col_label in enumerate(['#', 'English', 'Kanji', 'Hiragana']):
        c = atbl.rows[0].cells[i]
        set_cell_bg(c, NAVY)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(2)
        add_run(p, col_label, bold=True, size=10, color=WHITE)

    for row_i, (eng, kanji, hira, _ro, note) in enumerate(NUMBERS):
        row = atbl.add_row()
        bg = LIGHT if row_i % 2 == 0 else WHITE
        for ci, (val, clr, sz, bld) in enumerate([
            (str(row_i + 1), GREY,  9,  False),
            (eng,            DARK,  10, False),
            (kanji,          TEAL,  13, True),
            (hira,           NAVY,  10, False),
        ]):
            c = row.cells[ci]
            set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(2)
            if ci == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, val, bold=bld, size=sz, color=clr)
        if note:
            p_note = row.cells[3].paragraphs[0]
            add_run(p_note, f'\n★ {note}', italic=True, size=8, color=GREY)


def build_page6(doc):
    page_break(doc)
    add_page_label(doc, 6, TOTAL_PAGES, 'Answer Key — Teacher Copy')

    section_banner(doc, 'Page 6 — Matching Answer Key', bg=TEAL)

    _, answer_lookup = get_matching_bank()
    mtbl = doc.add_table(rows=1, cols=5)
    mtbl.style = 'Table Grid'
    for i, col_label in enumerate(['#', 'Kanji', 'Hiragana', 'Letter', 'English']):
        c = mtbl.rows[0].cells[i]
        set_cell_bg(c, TEAL)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(2)
        add_run(p, col_label, bold=True, size=10, color=WHITE)

    for row_i, (eng, kanji, hira, _ro, _note) in enumerate(MATCH_NUMBERS):
        row = mtbl.add_row()
        bg = MINT if row_i % 2 == 0 else WHITE
        eng_display = eng.split('—')[-1].strip()
        for ci, (val, clr, sz, bld) in enumerate([
            (str(row_i + 1),     GREY,  9,  True),
            (kanji,              TEAL, 14, True),
            (hira,               NAVY, 10, False),
            (answer_lookup[eng], TEAL, 10, True),
            (eng_display,        DARK, 10, False),
        ]):
            c = row.cells[ci]
            set_cell_bg(c, bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(3)
            add_run(p, val, bold=bld, size=sz, color=clr)


def build_page7(doc):
    page_break(doc)
    add_page_label(doc, 7, TOTAL_PAGES, 'Teacher Notes')

    section_banner(doc, "Page 7 — Teacher Notes & Activities", bg=NAVY)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(1)
    intro.paragraph_format.space_after = Pt(2)
    add_run(intro, 'Use this page as the teaching guide for the student worksheet.', size=10)

    notes_tbl = doc.add_table(rows=3, cols=2)
    notes_tbl.style = 'Table Grid'
    note_items = [
        ('Counting drill', 'Count 1-20 aloud as a class, then backwards in pairs.'),
        ('Kanji dictation', 'Say the English number and have students write only the kanji page first.'),
        ('Hiragana check', 'Repeat the same numbers and have students fill the hiragana page separately.'),
        ('Pattern discovery', 'Write 11-19 on the board and ask students what 十 is doing in each number.'),
        ('Matching review', 'Partners compare answer letters and explain how they matched each item.'),
        ('Extension', 'After 20, introduce 二十一, 三十五, 百, and 千 as pattern extensions.'),
    ]
    for idx, (title, desc) in enumerate(note_items):
        row_idx = idx // 2
        col_idx = idx % 2
        cell = notes_tbl.rows[row_idx].cells[col_idx]
        set_cell_bg(cell, LIGHT if row_idx % 2 == 0 else WHITE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(1)
        add_run(p, title, bold=True, size=11, color=NAVY)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        add_run(p2, desc, size=9, color=DARK)


# ── Entry point ────────────────────────────────────────────────────────────────
def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.2)
        section.bottom_margin = Inches(0.2)
        section.left_margin   = Inches(0.2)
        section.right_margin  = Inches(0.2)

    build_page1(doc)
    build_page2(doc)
    build_page3(doc)
    build_page4(doc)
    build_page5(doc)
    build_page6(doc)
    build_page7(doc)

    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'japanese_numbers_worksheet.docx')
    doc.save(out)
    print(f'Saved: {out}')


if __name__ == '__main__':
    main()
