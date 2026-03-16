from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins (tighter for better fit) ──────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x3A, 0x5C)
TEAL   = RGBColor(0x00, 0x7A, 0x87)
GOLD   = RGBColor(0xE6, 0xA8, 0x17)
LIGHT  = RGBColor(0xF0, 0xF6, 0xFA)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
GREY   = RGBColor(0x6B, 0x7B, 0x8D)

def _rgb_hex(rgb):
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  _rgb_hex(rgb))
    tcPr.append(shd)

def set_para_bg(para, rgb):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  _rgb_hex(rgb))
    pPr.append(shd)

def add_run(para, text, bold=False, italic=False, size=12, color=DARK, font='Calibri'):
    r = para.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = font
    return r

def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6 if level == 1 else 4)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, text, bold=True, size=14 if level == 1 else 12, color=NAVY if level == 1 else TEAL)
    return p

def body(text, indent=False, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    add_run(p, text, size=11)
    return p

def blank_line(text, label, hint='', indent_cm=0.6, space_after=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(indent_cm)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    add_run(p, f"{label}  ", bold=True, size=11, color=TEAL)
    add_run(p, text, size=11)
    if hint:
        add_run(p, f"  ({hint})", italic=True, size=10, color=GREY)
    return p

def page_break():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '007A87')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 1 — STUDENT WORKSHEET (Banner, Info, Objectives, Instructions, Part 1)
# ══════════════════════════════════════════════════════════════════════════════

tbl = doc.add_table(rows=2, cols=1)
tbl.style = 'Table Grid'
r0 = tbl.rows[0].cells[0]
set_cell_bg(r0, NAVY)
r0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p0 = r0.paragraphs[0]
p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
p0.paragraph_format.space_before, p0.paragraph_format.space_after = Pt(5), Pt(2)
add_run(p0, 'Beginner Japanese', bold=True, size=20, color=WHITE)
add_run(p0, '\nGreetings & Essential Phrases', bold=False, size=13, color=GOLD)
r1 = tbl.rows[1].cells[0]
set_cell_bg(r1, TEAL)
p1 = r1.paragraphs[0]
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.space_before, p1.paragraph_format.space_after = Pt(3), Pt(3)
add_run(p1, 'Student Worksheet  ·  Page 1 of 3', bold=False, size=10, color=WHITE)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

info = doc.add_table(rows=1, cols=4)
info.style = 'Table Grid'
for i, (lbl, blank) in enumerate([('Name', '_________________________________'), ('Class', '___________'), ('Date', '__________________'), ('Period', '____')]):
    c = info.rows[0].cells[i]
    set_cell_bg(c, LIGHT)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c.paragraphs[0]
    p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(2)
    add_run(p, f'{lbl}: ', bold=True, size=10, color=NAVY)
    add_run(p, blank, bold=False, size=10, color=DARK)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

heading('Learning Objectives')
for obj in ['Recognize and write basic Japanese greetings (hiragana).', 'Practice pronunciation using romaji hints.', 'Use phrases in short roleplay dialogues.']:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent, p.paragraph_format.space_before, p.paragraph_format.space_after = Cm(0.6), Pt(0), Pt(1)
    add_run(p, obj, size=10)
heading('Instructions')
body('For each English prompt below:')
for n, txt in [('1.', 'Write the Japanese (hiragana/kanji) on the first line.'), ('2.', 'Write the romaji on the second line.  Romaji hints are in parentheses.')]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent, p.paragraph_format.space_before, p.paragraph_format.space_after = Cm(0.6), Pt(0), Pt(1)
    add_run(p, f'{n} ', bold=True, size=10, color=TEAL)
    add_run(p, txt, size=10)
divider()

# Part 1 — all questions on this page (compact spacing)
p = doc.add_paragraph()
p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.left_indent = Pt(3), Pt(2), Cm(0.3)
add_run(p, 'Part 1 — Greetings & Phrases', bold=True, size=12, color=WHITE)
set_para_bg(p, NAVY)

questions = [
    ("Greeting someone in the morning",               "_" * 46, "_" * 32, "ohayou"),
    ("Greeting someone in the afternoon / daytime",   "_" * 46, "_" * 32, "konnichiwa"),
    ("Greeting someone in the evening / night",       "_" * 46, "_" * 32, "konbanwa"),
    ("Saying hello when meeting someone",             "_" * 46, "_" * 32, "konnichiwa"),
    ("Saying goodbye",                                "_" * 46, "_" * 32, "sayounara"),
    ("Asking someone 'Excuse me' before a question",  "_" * 46, "_" * 32, "sumimasen"),
    ("Asking someone for directions",                 "_" * 46, "_" * 32, "sumimasen, michi o oshiete kudasai"),
    ("Saying thank you",                              "_" * 46, "_" * 32, "arigatou"),
    ("Saying 'thank you very much'",                  "_" * 46, "_" * 32, "arigatou gozaimasu"),
    ("Saying sorry / excuse me",                      "_" * 46, "_" * 32, "gomen nasai"),
]
for i, (prompt, jp_blank, ro_blank, hint) in enumerate(questions, 1):
    pq = doc.add_paragraph()
    pq.paragraph_format.space_before, pq.paragraph_format.space_after = Pt(2), Pt(0)
    add_run(pq, f'{i}.  ', bold=True, size=10, color=TEAL)
    add_run(pq, prompt, bold=False, size=10, color=DARK)
    blank_line(jp_blank, 'Japanese:', hint='', indent_cm=0.8, space_after=0)
    blank_line(ro_blank, 'Romaji:  ', hint=hint, indent_cm=0.8, space_after=0)

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 2 — Part 2 Roleplay + Pronunciation Tips
# ══════════════════════════════════════════════════════════════════════════════
page_break()

p2_banner = doc.add_paragraph()
p2_banner.paragraph_format.space_before, p2_banner.paragraph_format.space_after = Pt(0), Pt(2)
add_run(p2_banner, 'Student Worksheet  ·  Page 2 of 3', bold=False, size=10, color=GREY)
p2_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.paragraph_format.space_before, p.paragraph_format.space_after, p.paragraph_format.left_indent = Pt(2), Pt(2), Cm(0.3)
add_run(p, 'Part 2 — Short Roleplay Practice  (pair up)', bold=True, size=12, color=WHITE)
set_para_bg(p, TEAL)
scenarios = [
    ('Scenario A — Meeting & Asking Directions', [('A', 'greeting'), ('B', 'response'), ('A', 'ask for directions'), ('B', 'short direction')]),
    ('Scenario B — Help & Thanks', [('A', 'asks for help'), ('B', 'helps'), ('A', 'says thank you'), ('B', 'polite reply')]),
]
for title, lines in scenarios:
    ps = doc.add_paragraph()
    ps.paragraph_format.space_before, ps.paragraph_format.space_after = Pt(4), Pt(1)
    add_run(ps, title, bold=True, size=10, color=NAVY)
    for speaker, cue in lines:
        pl = doc.add_paragraph()
        pl.paragraph_format.left_indent, pl.paragraph_format.space_before, pl.paragraph_format.space_after = Cm(0.8), Pt(1), Pt(1)
        add_run(pl, f'{speaker}:  ', bold=True, size=10, color=TEAL)
        add_run(pl, '_' * 38 + '  ', size=10)
        add_run(pl, f'({cue})', italic=True, size=9, color=GREY)
divider()

heading('Pronunciation Tips')
for tip in ['Vowels: a / e / i / o / u — short and steady.', 'ん (n) is nasal — hold it slightly.', 'Repeat slowly 3×, then 2× at normal speed.']:
    pt = doc.add_paragraph()
    pt.paragraph_format.left_indent, pt.paragraph_format.space_before, pt.paragraph_format.space_after = Cm(0.6), Pt(0), Pt(1)
    add_run(pt, '▸  ', bold=True, size=10, color=GOLD)
    add_run(pt, tip, size=10)

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE 3 — ANSWER KEY (Teacher)
# ══════════════════════════════════════════════════════════════════════════════
page_break()

tbl2 = doc.add_table(rows=2, cols=1)
tbl2.style = 'Table Grid'
ra0, ra1 = tbl2.rows[0].cells[0], tbl2.rows[1].cells[0]
set_cell_bg(ra0, DARK)
ra0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
pa0 = ra0.paragraphs[0]
pa0.alignment = WD_ALIGN_PARAGRAPH.CENTER
pa0.paragraph_format.space_before, pa0.paragraph_format.space_after = Pt(5), Pt(2)
add_run(pa0, 'Answer Key — Teacher Copy', bold=True, size=16, color=WHITE)
add_run(pa0, '\nPage 3 of 3', bold=False, size=11, color=GOLD)
set_cell_bg(ra1, GOLD)
pa1 = ra1.paragraphs[0]
pa1.alignment = WD_ALIGN_PARAGRAPH.CENTER
pa1.paragraph_format.space_before, pa1.paragraph_format.space_after = Pt(2), Pt(2)
add_run(pa1, 'For teacher use only — do not distribute to students', bold=False, size=9, color=DARK)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

answers = [
    ("1. Morning greeting",          "おはよう",                          "ohayou",                               "Casual. Polite: おはようございます (ohayou gozaimasu)."),
    ("2. Afternoon greeting",        "こんにちは",                        "konnichiwa",                           "General daytime greeting."),
    ("3. Evening greeting",          "こんばんは",                        "konbanwa",                             "Used after sunset / in the evening."),
    ("4. Hello when meeting",        "こんにちは",                        "konnichiwa",                           "Same as #2 for daytime meetings."),
    ("5. Goodbye",                   "さようなら",                        "sayounara",                            "Formal; used when parting."),
    ("6. Excuse me (before Q)",      "すみません",                        "sumimasen",                            "Also used to get attention or lightly apologize."),
    ("7. Asking for directions",     "すみません、みちをおしえてください", "sumimasen, michi o oshiete kudasai",   "Polite request: 'Please tell me the way.'"),
    ("8. Thank you",                 "ありがとう",                        "arigatou",                             "Casual. Polite: ありがとうございます (arigatou gozaimasu)."),
    ("9. Thank you very much",       "ありがとうございます",              "arigatou gozaimasu",                   "Polite; appropriate in most situations."),
    ("10. Sorry / excuse me",        "ごめんなさい",                      "gomen nasai",                          "Used to apologize; casual to neutral tone."),
]
atbl = doc.add_table(rows=1, cols=4)
atbl.style = 'Table Grid'
for i, txt in enumerate(['Prompt', 'Japanese', 'Romaji', 'Notes']):
    c = atbl.rows[0].cells[i]
    set_cell_bg(c, NAVY)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(2)
    add_run(p, txt, bold=True, size=10, color=WHITE)
for row_i, (prompt, jp, ro, note) in enumerate(answers):
    row = atbl.add_row()
    bg = LIGHT if row_i % 2 == 0 else WHITE
    for ci, val in enumerate([prompt, jp, ro, note]):
        c = row.cells[ci]
        set_cell_bg(c, bg)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(2)
        color = TEAL if ci == 1 else (GREY if ci == 3 else DARK)
        add_run(p, val, bold=(ci == 1), size=10 if ci != 3 else 9, color=color)
doc.add_paragraph().paragraph_format.space_after = Pt(4)
divider()

heading("Teacher's Notes & Suggested Activities")
for title, desc in [('Pronunciation drill (5 min)', 'Teacher says each phrase; class repeats 3×.'), ('Pair roleplay (10–12 min)', 'Students practice Parts A/B and swap roles.'), ('Quick written quiz (5 min)', 'Teacher reads English prompts; students write Japanese.'), ('Homework extension', 'Students submit two polite phrases, e.g. はじめまして (hajimemashite).')]:
    pa = doc.add_paragraph()
    pa.paragraph_format.left_indent, pa.paragraph_format.space_before, pa.paragraph_format.space_after = Cm(0.6), Pt(1), Pt(1)
    add_run(pa, f'{title}: ', bold=True, size=10, color=NAVY)
    add_run(pa, desc, bold=False, size=10, color=DARK)
divider()

heading('Printing & Formatting Tips')
for t in ['Print double-sided: student page front, answer key back.', 'Increase line spacing to 1.3–1.5 for extra handwriting space.', 'Font: Calibri or Arial; Title 20–24 pt bold; Body 12 pt.']:
    pf = doc.add_paragraph()
    pf.paragraph_format.left_indent, pf.paragraph_format.space_before, pf.paragraph_format.space_after = Cm(0.6), Pt(0), Pt(1)
    add_run(pf, '▸  ', bold=True, size=10, color=GOLD)
    add_run(pf, t, size=10)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'japanese_greetings_worksheet.docx')
doc.save(out)
print(f'Saved: {out}')
