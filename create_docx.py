"""Create a DOCX document using python-docx."""
from docx import Document
from docx.shared import Pt

doc = Document()
doc.add_heading("My Document", 0)
doc.add_paragraph("Hello. This is a paragraph.")
doc.add_heading("Section", level=1)
p = doc.add_paragraph()
p.add_run("Bold text. ").bold = True
p.add_run("Normal text.")
doc.save("output.docx")
print("Saved output.docx")
